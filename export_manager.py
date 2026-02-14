"""
Export Manager Module for MarkItDown Notepad

Provides export functionality to PDF and DOCX formats.
Uses 'markdown' for HTML generation, 'weasyprint' for PDF rendering,
and 'python-docx' for Word document creation.
"""

import os
import re
import base64
import tempfile
from pathlib import Path
from typing import Optional, Dict, Tuple

# Try to import markdown
try:
    import markdown
    from markdown.extensions.tables import TableExtension
    from markdown.extensions.fenced_code import FencedCodeExtension
    from markdown.extensions.codehilite import CodeHiliteExtension
    from markdown.extensions.toc import TocExtension
    MARKDOWN_AVAILABLE = True
except ImportError:
    MARKDOWN_AVAILABLE = False

# Try to import weasyprint for PDF
try:
    from weasyprint import HTML as WeasyprintHTML
    WEASYPRINT_AVAILABLE = True
except ImportError:
    WEASYPRINT_AVAILABLE = False

# Try to import python-docx for DOCX
try:
    from docx import Document as DocxDocument
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Try to import PIL for image handling in exports
try:
    from PIL import Image
    PIL_AVAILABLE = True
except ImportError:
    PIL_AVAILABLE = False


def get_available_export_formats() -> Dict[str, bool]:
    """Return dict of format name -> availability"""
    return {
        'pdf': WEASYPRINT_AVAILABLE and MARKDOWN_AVAILABLE,
        'docx': DOCX_AVAILABLE,
        'html': MARKDOWN_AVAILABLE,
    }


# =============================================================================
# CSS Stylesheet for PDF/HTML Export
# =============================================================================

def _build_css(theme: Optional[Dict] = None, page_settings: Optional[Dict] = None) -> str:
    """Build a CSS stylesheet for HTML/PDF export based on theme"""
    # Defaults
    bg = theme.get('visual_bg', '#ffffff') if theme else '#ffffff'
    fg = theme.get('visual_fg', '#1a1a1a') if theme else '#1a1a1a'
    font_family = theme.get('visual_font', 'Georgia') if theme else 'Georgia'
    font_size = theme.get('visual_font_size', 11) if theme else 11
    
    page_size = 'A4'
    margin = '2cm'
    if page_settings:
        page_size = page_settings.get('page_size', 'A4')
        margin = page_settings.get('margin', '2cm')
    
    return f"""
    @page {{
        size: {page_size};
        margin: {margin};
    }}
    body {{
        font-family: '{font_family}', Georgia, 'Times New Roman', serif;
        font-size: {font_size}pt;
        line-height: 1.6;
        color: {fg};
        background-color: {bg};
        max-width: 100%;
        padding: 0;
    }}
    h1 {{ font-size: 2em; margin: 0.8em 0 0.4em; border-bottom: 2px solid #ccc; padding-bottom: 0.3em; }}
    h2 {{ font-size: 1.6em; margin: 0.7em 0 0.35em; border-bottom: 1px solid #ddd; padding-bottom: 0.2em; }}
    h3 {{ font-size: 1.3em; margin: 0.6em 0 0.3em; }}
    h4 {{ font-size: 1.1em; margin: 0.5em 0 0.25em; }}
    h5, h6 {{ font-size: 1em; margin: 0.5em 0 0.25em; }}
    p {{ margin: 0.5em 0; }}
    blockquote {{
        border-left: 4px solid #ccc;
        margin: 1em 0;
        padding: 0.5em 1em;
        color: #555;
        background: #f9f9f9;
    }}
    code {{
        font-family: 'Consolas', 'Monaco', 'Courier New', monospace;
        background: #f4f4f4;
        padding: 0.15em 0.4em;
        border-radius: 3px;
        font-size: 0.9em;
    }}
    pre {{
        background: #f4f4f4;
        padding: 1em;
        border-radius: 5px;
        overflow-x: auto;
        border: 1px solid #ddd;
    }}
    pre code {{
        background: none;
        padding: 0;
    }}
    table {{
        border-collapse: collapse;
        width: 100%;
        margin: 1em 0;
    }}
    th, td {{
        border: 1px solid #ccc;
        padding: 0.5em 0.75em;
        text-align: left;
    }}
    th {{
        background: #f0f0f0;
        font-weight: bold;
    }}
    tr:nth-child(even) {{
        background: #fafafa;
    }}
    img {{
        max-width: 100%;
        height: auto;
    }}
    hr {{
        border: none;
        border-top: 1px solid #ccc;
        margin: 1.5em 0;
    }}
    ul, ol {{
        margin: 0.5em 0;
        padding-left: 2em;
    }}
    li {{
        margin: 0.2em 0;
    }}
    a {{
        color: #0066cc;
        text-decoration: none;
    }}
    a:hover {{
        text-decoration: underline;
    }}
    """


# =============================================================================
# Markdown to HTML Conversion
# =============================================================================

def _markdown_to_html(md_content: str, theme: Optional[Dict] = None,
                       page_settings: Optional[Dict] = None,
                       base_path: Optional[str] = None) -> str:
    """Convert markdown content to a full HTML document"""
    if not MARKDOWN_AVAILABLE:
        raise RuntimeError("Python 'markdown' package is required. Install with: pip install markdown")
    
    extensions = [
        'tables',
        'fenced_code',
        'codehilite',
        'toc',
        'nl2br',
        'sane_lists',
        'smarty',
    ]
    
    extension_configs = {
        'codehilite': {
            'css_class': 'highlight',
            'guess_lang': False,
        },
    }
    
    html_body = markdown.markdown(md_content, extensions=extensions,
                                   extension_configs=extension_configs)
    
    # Resolve relative image paths to absolute for PDF rendering
    if base_path:
        html_body = _resolve_image_paths(html_body, base_path)
    
    # Handle base64 images (already work as data URIs)
    css = _build_css(theme, page_settings)
    
    html = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <style>{css}</style>
</head>
<body>
{html_body}
</body>
</html>"""
    
    return html


def _resolve_image_paths(html: str, base_path: str) -> str:
    """Resolve relative image paths in HTML to absolute file:// URIs"""
    import re
    
    def replace_src(match):
        src = match.group(1)
        # Skip data URIs and absolute URLs
        if src.startswith(('data:', 'http://', 'https://', 'file://')):
            return match.group(0)
        # Resolve relative path
        abs_path = os.path.abspath(os.path.join(base_path, src))
        if os.path.exists(abs_path):
            return f'src="file://{abs_path}"'
        return match.group(0)
    
    return re.sub(r'src="([^"]*)"', replace_src, html)


# =============================================================================
# Export to PDF
# =============================================================================

def export_to_pdf(md_content: str, output_path: str,
                   theme: Optional[Dict] = None,
                   page_settings: Optional[Dict] = None,
                   base_path: Optional[str] = None) -> bool:
    """
    Export markdown content to PDF.
    
    Args:
        md_content: Markdown text
        output_path: Path for the output PDF file
        theme: Theme dict with font/color settings
        page_settings: Dict with 'page_size' and 'margin'
        base_path: Base directory for resolving relative image paths
    
    Returns:
        True on success
    
    Raises:
        RuntimeError: If required libraries are not installed
    """
    if not WEASYPRINT_AVAILABLE:
        raise RuntimeError("WeasyPrint is required for PDF export. Install with: pip install weasyprint")
    if not MARKDOWN_AVAILABLE:
        raise RuntimeError("Python 'markdown' package is required. Install with: pip install markdown")
    
    html = _markdown_to_html(md_content, theme, page_settings, base_path)
    
    WeasyprintHTML(string=html, base_url=base_path or '.').write_pdf(output_path)
    return True


# =============================================================================
# Export to HTML
# =============================================================================

def export_to_html(md_content: str, output_path: str,
                    theme: Optional[Dict] = None,
                    base_path: Optional[str] = None) -> bool:
    """
    Export markdown content to standalone HTML.
    
    Args:
        md_content: Markdown text
        output_path: Path for the output HTML file
        theme: Theme dict
        base_path: Base directory for resolving relative image paths
    
    Returns:
        True on success
    """
    if not MARKDOWN_AVAILABLE:
        raise RuntimeError("Python 'markdown' package is required. Install with: pip install markdown")
    
    html = _markdown_to_html(md_content, theme, base_path=base_path)
    
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(html)
    return True


# =============================================================================
# Export to DOCX
# =============================================================================

def export_to_docx(md_content: str, output_path: str,
                    theme: Optional[Dict] = None,
                    base_path: Optional[str] = None) -> bool:
    """
    Export markdown content to DOCX (Word document).
    
    Parses the markdown into structural elements and builds a python-docx Document
    with heading styles, paragraph formatting, code blocks, lists, tables, and images.
    
    Args:
        md_content: Markdown text
        output_path: Path for the output DOCX file
        theme: Theme dict
        base_path: Base directory for resolving relative image paths
    
    Returns:
        True on success
    """
    if not DOCX_AVAILABLE:
        raise RuntimeError("python-docx is required for DOCX export. Install with: pip install python-docx")
    
    doc = DocxDocument()
    
    # Configure default style
    style = doc.styles['Normal']
    font_name = theme.get('visual_font', 'Calibri') if theme else 'Calibri'
    font_size = theme.get('visual_font_size', 11) if theme else 11
    style.font.name = font_name
    style.font.size = Pt(font_size)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15
    
    # Setup code style
    _ensure_code_style(doc)
    
    lines = md_content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i]
        
        # Fenced code block
        if line.strip().startswith('```'):
            code_lines = []
            lang = line.strip()[3:].strip()
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            i += 1  # skip closing ```
            _add_code_block(doc, '\n'.join(code_lines), lang)
            continue
        
        # Heading
        heading_match = re.match(r'^(#{1,6})\s+(.*)', line)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            doc.add_heading(text, level=min(level, 9))
            i += 1
            continue
        
        # Horizontal rule
        if re.match(r'^[\s]*([-*_])\s*\1\s*\1[\s\-\*_]*$', line):
            # Add a thin paragraph as a visual separator
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run('─' * 50)
            run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
            run.font.size = Pt(8)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue
        
        # Blockquote
        if line.startswith('>'):
            quote_lines = []
            while i < len(lines) and lines[i].startswith('>'):
                quote_lines.append(lines[i].lstrip('>').strip())
                i += 1
            p = doc.add_paragraph(' '.join(quote_lines))
            p.style = doc.styles['Normal']
            p.paragraph_format.left_indent = Cm(1.5)
            pf = p.paragraph_format
            pf.space_before = Pt(4)
            pf.space_after = Pt(4)
            for run in p.runs:
                run.font.italic = True
                run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            continue
        
        # Unordered list
        if re.match(r'^[\s]*[-*+]\s+', line):
            while i < len(lines) and re.match(r'^[\s]*[-*+]\s+', lines[i]):
                text = re.sub(r'^[\s]*[-*+]\s+', '', lines[i])
                p = doc.add_paragraph(style='List Bullet')
                _add_inline_formatting(p, text)
                i += 1
            continue
        
        # Ordered list
        if re.match(r'^[\s]*\d+\.\s+', line):
            while i < len(lines) and re.match(r'^[\s]*\d+\.\s+', lines[i]):
                text = re.sub(r'^[\s]*\d+\.\s+', '', lines[i])
                p = doc.add_paragraph(style='List Number')
                _add_inline_formatting(p, text)
                i += 1
            continue
        
        # Table
        if '|' in line and i + 1 < len(lines) and re.match(r'^[\s]*\|[\s:]*[-]+', lines[i + 1]):
            table_lines = []
            while i < len(lines) and '|' in lines[i]:
                table_lines.append(lines[i])
                i += 1
            _add_table(doc, table_lines)
            continue
        
        # Image
        img_match = re.match(r'!\[([^\]]*)\]\(([^)]+)\)', line.strip())
        if img_match:
            alt_text = img_match.group(1)
            img_path = img_match.group(2)
            _add_image(doc, img_path, alt_text, base_path)
            i += 1
            continue
        
        # Blank line
        if not line.strip():
            i += 1
            continue
        
        # Regular paragraph — collect contiguous non-blank, non-special lines
        para_lines = []
        while i < len(lines) and lines[i].strip() and not _is_special_line(lines[i]):
            para_lines.append(lines[i])
            i += 1
        
        if para_lines:
            p = doc.add_paragraph()
            _add_inline_formatting(p, ' '.join(para_lines))
    
    doc.save(output_path)
    return True


def _is_special_line(line: str) -> bool:
    """Check if a line is a special markdown element"""
    stripped = line.strip()
    if stripped.startswith('#'):
        return True
    if stripped.startswith('```'):
        return True
    if stripped.startswith('>'):
        return True
    if re.match(r'^[\s]*[-*+]\s+', line):
        return True
    if re.match(r'^[\s]*\d+\.\s+', line):
        return True
    if re.match(r'^[\s]*([-*_])\s*\1\s*\1[\s\-\*_]*$', line):
        return True
    if re.match(r'^!\[', stripped):
        return True
    if '|' in line and re.match(r'^\s*\|', line):
        return True
    return False


def _ensure_code_style(doc):
    """Ensure a 'Code' paragraph style exists in the document"""
    try:
        doc.styles['Code']
    except KeyError:
        style = doc.styles.add_style('Code', WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = 'Consolas'
        style.font.size = Pt(9)
        style.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
        style.paragraph_format.space_before = Pt(4)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.left_indent = Cm(0.5)


def _add_code_block(doc, code: str, language: str = ""):
    """Add a code block to the document"""
    p = doc.add_paragraph(style='Code')
    run = p.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    # Add background shading via XML (light gray)
    from docx.oxml.ns import qn
    shading = p._element.get_or_add_pPr()
    shd = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): 'F4F4F4'
    })
    shading.append(shd)


def _add_inline_formatting(paragraph, text: str):
    """Parse inline markdown formatting and add runs to a paragraph"""
    # Pattern for **bold**, *italic*, `code`, [links](url)
    pattern = re.compile(
        r'(\*\*(.+?)\*\*)'       # bold
        r'|(\*(.+?)\*)'            # italic
        r'|(`(.+?)`)'              # inline code
        r'|(\[([^\]]+)\]\(([^)]+)\))'  # link
        r'|(~~(.+?)~~)'            # strikethrough
    )
    
    last_end = 0
    for match in pattern.finditer(text):
        # Add text before match
        if match.start() > last_end:
            paragraph.add_run(text[last_end:match.start()])
        
        if match.group(2):  # bold
            run = paragraph.add_run(match.group(2))
            run.bold = True
        elif match.group(4):  # italic
            run = paragraph.add_run(match.group(4))
            run.italic = True
        elif match.group(6):  # code
            run = paragraph.add_run(match.group(6))
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
        elif match.group(8):  # link
            link_text = match.group(8)
            link_url = match.group(9)
            run = paragraph.add_run(f"{link_text} ({link_url})")
            run.font.color.rgb = RGBColor(0x00, 0x66, 0xCC)
            run.underline = True
        elif match.group(11):  # strikethrough
            run = paragraph.add_run(match.group(11))
            run.font.strike = True
        
        last_end = match.end()
    
    # Add remaining text
    if last_end < len(text):
        paragraph.add_run(text[last_end:])


def _add_table(doc, table_lines: list):
    """Add a markdown table to the document"""
    if len(table_lines) < 2:
        return
    
    # Parse header
    headers = [c.strip() for c in table_lines[0].split('|') if c.strip()]
    
    # Parse data rows (skip separator line)
    rows = []
    for line in table_lines[2:]:
        cells = [c.strip() for c in line.split('|') if c.strip()]
        if cells:
            rows.append(cells)
    
    if not headers:
        return
    
    num_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=num_cols)
    table.style = 'Table Grid'
    
    # Header row
    for j, header in enumerate(headers):
        if j < num_cols:
            cell = table.rows[0].cells[j]
            cell.text = header
            for p in cell.paragraphs:
                for run in p.runs:
                    run.bold = True
    
    # Data rows
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            if j < num_cols:
                table.rows[i + 1].cells[j].text = cell_text
    
    # Add spacing after table
    doc.add_paragraph()


def _add_image(doc, img_path: str, alt_text: str, base_path: Optional[str] = None):
    """Add an image to the document"""
    # Handle base64 data URIs
    if img_path.startswith('data:'):
        _add_base64_image(doc, img_path, alt_text)
        return
    
    # Resolve relative path
    if base_path and not os.path.isabs(img_path):
        img_path = os.path.abspath(os.path.join(base_path, img_path))
    
    if os.path.exists(img_path):
        try:
            # Determine width — fit to page (max ~6 inches)
            doc.add_picture(img_path, width=Inches(5.5))
            if alt_text:
                p = doc.add_paragraph(alt_text)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.italic = True
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        except Exception as e:
            doc.add_paragraph(f"[Image: {alt_text or img_path}] (Failed to embed: {e})")
    else:
        doc.add_paragraph(f"[Image: {alt_text or img_path}] (File not found)")


def _add_base64_image(doc, data_uri: str, alt_text: str):
    """Add a base64-encoded image to the document"""
    try:
        # Parse data URI: data:image/png;base64,xxxxx
        header, data = data_uri.split(',', 1)
        img_bytes = base64.b64decode(data)
        
        # Write to temp file
        suffix = '.png'
        if 'jpeg' in header or 'jpg' in header:
            suffix = '.jpg'
        elif 'gif' in header:
            suffix = '.gif'
        
        with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp:
            tmp.write(img_bytes)
            tmp_path = tmp.name
        
        try:
            doc.add_picture(tmp_path, width=Inches(5.5))
            if alt_text:
                p = doc.add_paragraph(alt_text)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.italic = True
                    run.font.size = Pt(9)
        finally:
            os.unlink(tmp_path)
    except Exception as e:
        doc.add_paragraph(f"[Image: {alt_text}] (Failed to decode base64: {e})")
